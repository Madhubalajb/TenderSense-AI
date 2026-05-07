"""
extractor.py — Document Text Extractor (No Java, No Tika)
-----------------------------------------------------------
Reads any document format and returns plain text.

Libraries used:
  PyMuPDF  (fitz)  — PDFs, no Java needed, very fast
  python-docx      — Word documents (.docx)
  openpyxl         — Excel files (.xlsx)
  Pillow           — images (.jpg, .png, etc.)
  OCR pipeline     — for scanned PDFs and images

Install:
  pip install pymupdf python-docx openpyxl pillow
"""

import os
from pathlib import Path
from app.pipeline.ocr import ocr_file

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"}


def extract_text(file_path: str) -> dict:
    """
    Extract text from any supported file.

    Returns:
        {
          text:    extracted plain text (str),
          method:  how it was extracted (str),
          page_count: number of pages (int or None),
          warnings: list of warning strings,
          error:   error message or None,
          file_name: original file name,
        }
    """
    result = {
        "text":       "",
        "method":     None,
        "page_count": None,
        "warnings":   [],
        "error":      None,
        "file_name":  Path(file_path).name,
    }

    if not os.path.exists(file_path):
        result["error"] = f"File not found: {file_path}"
        return result

    ext = Path(file_path).suffix.lower()

    # ── Images — always OCR ────────────────────────────────────────────────────
    if ext in IMAGE_EXTENSIONS:
        return _do_ocr(file_path, result)

    # ── PDF ────────────────────────────────────────────────────────────────────
    if ext == ".pdf":
        return _extract_pdf(file_path, result)

    # ── Word document ──────────────────────────────────────────────────────────
    if ext in {".docx", ".doc"}:
        return _extract_docx(file_path, result)

    # ── Excel ──────────────────────────────────────────────────────────────────
    if ext in {".xlsx", ".xls"}:
        return _extract_excel(file_path, result)

    # ── Plain text ─────────────────────────────────────────────────────────────
    if ext in {".txt", ".md"}:
        return _extract_txt(file_path, result)

    result["error"] = f"Unsupported file type: {ext}"
    return result


# ── PDF via PyMuPDF ────────────────────────────────────────────────────────────

def _extract_pdf(file_path: str, result: dict) -> dict:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        result["error"] = (
            "PyMuPDF not installed. Run: pip install pymupdf"
        )
        return result

    try:
        doc        = fitz.open(file_path)
        pages_text = []
        total_chars = 0

        for page_num in range(len(doc)):
            page = doc[page_num]

            # Use "blocks" extraction to get text with spatial layout preserved.
            # This gives us text in reading order with blank lines between blocks,
            # which is far better for tables than a flat "text" dump.
            blocks = page.get_text("blocks", sort=True)  # sorted top-to-bottom
            lines  = []
            prev_y = None
            for b in blocks:
                # b = (x0, y0, x1, y1, text, block_no, block_type)
                if b[6] != 0:   # skip non-text blocks (images)
                    continue
                block_text = b[4].strip()
                if not block_text:
                    continue
                # Insert a blank line when there's a large vertical gap
                # (heuristic: gap > 15pt signals a new section/row)
                if prev_y is not None and (b[1] - prev_y) > 15:
                    lines.append("")
                lines.append(block_text)
                prev_y = b[3]   # bottom y of current block

            text = "\n".join(lines)
            pages_text.append(text)
            total_chars += len(text.strip())

        result["page_count"] = len(doc)
        doc.close()

        # If very little text extracted — likely a scanned PDF
        if total_chars < 100:
            result["warnings"].append(
                "PDF appears to be scanned (little selectable text found). "
                "Switching to OCR — accuracy depends on scan quality."
            )
            return _do_ocr(file_path, result)

        full_text       = "\n\n".join(pages_text)
        result["text"]  = _clean(full_text)
        result["method"] = "pymupdf"

        print(f"[extractor] PyMuPDF: {result['page_count']} pages, "
              f"{len(result['text']):,} chars")
        return result

    except Exception as e:
        result["error"] = f"PDF extraction failed: {e}"
        return result


# ── Word document via python-docx ──────────────────────────────────────────────

def _extract_docx(file_path: str, result: dict) -> dict:
    try:
        from docx import Document
    except ImportError:
        result["error"] = "python-docx not installed. Run: pip install python-docx"
        return result

    try:
        doc   = Document(file_path)
        lines = [para.text for para in doc.paragraphs if para.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    lines.append(row_text)

        result["text"]   = _clean("\n".join(lines))
        result["method"] = "python-docx"
        print(f"[extractor] python-docx: {len(lines)} paragraphs")
        return result

    except Exception as e:
        result["error"] = f"DOCX extraction failed: {e}"
        return result


# ── Excel via openpyxl ─────────────────────────────────────────────────────────

def _extract_excel(file_path: str, result: dict) -> dict:
    try:
        import openpyxl
    except ImportError:
        result["error"] = "openpyxl not installed. Run: pip install openpyxl"
        return result

    try:
        wb    = openpyxl.load_workbook(file_path, data_only=True)
        lines = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(
                    str(cell) for cell in row
                    if cell is not None and str(cell).strip()
                )
                if row_text:
                    lines.append(row_text)

        result["text"]   = _clean("\n".join(lines))
        result["method"] = "openpyxl"
        print(f"[extractor] openpyxl: {len(wb.sheetnames)} sheets")
        return result

    except Exception as e:
        result["error"] = f"Excel extraction failed: {e}"
        return result


# ── Plain text ─────────────────────────────────────────────────────────────────

def _extract_txt(file_path: str, result: dict) -> dict:
    try:
        with open(file_path, encoding="utf-8", errors="replace") as f:
            result["text"]   = f.read()
        result["method"] = "plaintext"
        return result
    except Exception as e:
        result["error"] = f"Text file read failed: {e}"
        return result


# ── OCR fallback ───────────────────────────────────────────────────────────────

def _do_ocr(file_path: str, result: dict) -> dict:
    print(f"[extractor] Routing to OCR: {file_path}")
    ocr = ocr_file(file_path)
    result["text"]   = ocr["text"]
    result["method"] = "ocr"
    result["warnings"].extend(ocr.get("warnings", []))
    if ocr.get("confidence", 100) < 80:
        result["warnings"].append(
            f"Low OCR confidence ({ocr['confidence']:.0f}%). "
            "Document may need manual review."
        )
    return result


# ── Text cleaning ──────────────────────────────────────────────────────────────

def _clean(text: str) -> str:
    """Collapse excessive blank lines."""
    if not text:
        return ""
    lines   = text.splitlines()
    cleaned = []
    blanks  = 0
    for line in lines:
        if line.strip() == "":
            blanks += 1
            if blanks <= 2:
                cleaned.append("")
        else:
            blanks = 0
            cleaned.append(line.strip())
    return "\n".join(cleaned).strip()


# ── Quick test ─────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python -m app.pipeline.extractor <path-to-file>")
        sys.exit(1)

    path = sys.argv[1]
    print(f"\nExtracting: {path}\n{'─'*50}")
    r = extract_text(path)

    if r["error"]:
        print(f"ERROR: {r['error']}")
    else:
        print(f"Method     : {r['method']}")
        print(f"Pages      : {r['page_count'] or 'n/a'}")
        print(f"Characters : {len(r['text']):,}")
        print(f"Warnings   : {r['warnings'] or 'none'}")
        print(f"\n── First 1500 characters ──\n")
        print(r["text"][:1500])